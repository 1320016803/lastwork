from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.table import _Cell

from .converter_base import BaseConverter


class DocxConverter(BaseConverter):
    SUPPORTED_EXT = (".docx",)

    # 标题 style 名 → Markdown 标题等级
    HEADING_MAP = {
        "Heading 1": "#",
        "Heading 2": "##",
        "Heading 3": "###",
        "Heading 4": "####",
        "Heading 5": "#####",
        "Heading 6": "######",
        "标题 1": "#",
        "标题 2": "##",
        "标题 3": "###",
        "标题 4": "####",
        "标题 5": "#####",
        "标题 6": "######",
    }

    def _to_markdown(self, src: Path) -> str:
        doc = Document(str(src))

        # 用 XML 元素 id 做索引，确保段落 / 表格出现顺序与源文档一致
        p_map = {id(p._element): p for p in doc.paragraphs}
        t_map = {id(t._element): t for t in doc.tables}

        lines: list[str] = []
        for block in doc.element.body:
            bid = id(block)
            if bid in p_map:
                text = self._render_paragraph(p_map[bid])
                lines.append(text)
            elif bid in t_map:
                lines.append("")
                lines.extend(self._render_table(t_map[bid]))
                lines.append("")

        return "\n".join(self._collapse_blank_lines(lines)).strip() + "\n"

    def _render_paragraph(self, paragraph) -> str:
        style_name = paragraph.style.name if paragraph.style else "Normal"
        raw_text = self._render_inline_text(paragraph)

        if style_name in self.HEADING_MAP:
            return f"{self.HEADING_MAP[style_name]} {raw_text}"

        if style_name and ("Quote" in style_name or "引文" in style_name):
            return "> " + raw_text

        if style_name and ("List" in style_name or "列表" in style_name):
            return "- " + raw_text

        return raw_text

    def _render_inline_text(self, paragraph) -> str:
        """拼接段落内的文本片段，处理粗体、斜体等格式"""
        parts: list[str] = []
        for run in paragraph.runs:
            text = run.text
            if not text:
                continue
            if run.bold and run.italic:
                text = f"***{text}***"
            elif run.bold:
                text = f"**{text}**"
            elif run.italic:
                text = f"*{text}*"
            if run.underline and run.underline is not None and str(run.underline) != "None":
                text = f"<u>{text}</u>"
            parts.append(text)
        return "".join(parts)

    # ---------- 表格 ----------
    def _render_table(self, table: Table) -> list[str]:
        if not table.rows:
            return []

        rendered: list[list[str]] = []
        for row in table.rows:
            rendered.append([self._cell_text(cell) for cell in row.cells])

        n_cols = max(len(row) for row in rendered)
        normalized = [row + [""] * (n_cols - len(row)) for row in rendered]

        header_row = normalized[0]
        sep_row = ["---"] * n_cols
        md_lines = [
            "| " + " | ".join(cell.replace("\n", " ").replace("|", "\\|") for cell in header_row) + " |",
            "| " + " | ".join(sep_row) + " |",
        ]
        for row in normalized[1:]:
            md_lines.append(
                "| " + " | ".join(cell.replace("\n", " ").replace("|", "\\|") for cell in row) + " |"
            )
        return md_lines

    def _cell_text(self, cell: _Cell) -> str:
        return " ".join(p.text for p in cell.paragraphs).strip()

    # ---------- 工具 ----------
    def _collapse_blank_lines(self, lines: list[str]) -> list[str]:
        result: list[str] = []
        blank_count = 0
        for line in lines:
            if line.strip() == "":
                blank_count += 1
                if blank_count <= 2:
                    result.append("")
            else:
                blank_count = 0
                result.append(line)
        return result
