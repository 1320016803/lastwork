"""端到端转换测试：生成示例 docx/pdf/txt/xlsx 并转换为 md"""
import sys
from pathlib import Path

# 确保 Windows cmd 下的 UTF-8 输出
if sys.stdout.encoding and "UTF" not in sys.stdout.encoding.upper():
    sys.stdout.reconfigure(encoding="utf-8", errors="replace")
    sys.stderr.reconfigure(encoding="utf-8", errors="replace")

APP_DIR = Path(__file__).resolve().parent
sys.path.insert(0, str(APP_DIR))

from converters.converter_dispatcher import convert_file


def ensure_samples() -> dict:
    """生成示例文档，返回 {ext: path}"""
    samples_dir = APP_DIR / "_sample_inputs"
    samples_dir.mkdir(exist_ok=True)

    files = {}

    # 1. TXT
    txt_path = samples_dir / "示例文本.txt"
    txt_path.write_text(
        "这是一个示例文档。\n"
        "用来测试文档转 Markdown 的功能。\n\n"
        "第二段落的内容。\n",
        encoding="utf-8",
    )
    files["txt"] = str(txt_path)

    # 2. DOCX
    from docx import Document
    doc = Document()
    doc.add_heading("示例文档标题", level=1)
    doc.add_paragraph("这是第一个段落，包含一些普通文本内容。")
    doc.add_heading("二级标题", level=2)
    doc.add_paragraph("这里是第二个段落，列表如下：")
    doc.add_paragraph("第一项", style="List Bullet")
    doc.add_paragraph("第二项", style="List Bullet")
    table = doc.add_table(rows=2, cols=3)
    table.style = "Table Grid"
    for i, val in enumerate(["姓名", "年龄", "城市"]):
        table.rows[0].cells[i].text = val
    for i, val in enumerate(["张三", "25", "北京"]):
        table.rows[1].cells[i].text = val
    docx_path = samples_dir / "示例文档.docx"
    doc.save(str(docx_path))
    files["docx"] = str(docx_path)

    # 3. XLSX
    from openpyxl import Workbook
    wb = Workbook()
    ws = wb.active
    ws.title = "成员"
    ws.append(["姓名", "年龄", "城市", "分数"])
    ws.append(["张三", 25, "北京", 95.5])
    ws.append(["李四", 30, "上海", 88])
    ws.append(["王五", 27, "广州", 76])
    ws2 = wb.create_sheet("备注")
    ws2.append(["这是第二个工作表的内容"])
    ws2.append(["hello", "world"])
    xlsx_path = samples_dir / "示例表格.xlsx"
    wb.save(str(xlsx_path))
    files["xlsx"] = str(xlsx_path)

    # 4. PDF
    import fitz
    pdf_path = samples_dir / "示例PDF.pdf"
    pdf_doc = fitz.open()
    page = pdf_doc.new_page()
    page.insert_text((50, 50), "PDF 示例文档", fontsize=16)
    page.insert_text((50, 90), "这是从 PDF 提取的示例文本。", fontsize=11)
    page.insert_text((50, 120), "用于验证 pdf_converter 的功能。", fontsize=11)
    pdf_doc.save(str(pdf_path))
    pdf_doc.close()
    files["pdf"] = str(pdf_path)

    return files


def main() -> int:
    print(">>> 生成示例文档 ...")
    files = ensure_samples()
    for ext, path in files.items():
        print(f"    [{ext}] {path}")

    output_dir = APP_DIR / "_sample_outputs"
    output_dir.mkdir(exist_ok=True)

    print("\n>>> 开始转换 ...")
    errors: list[str] = []
    for ext, path in files.items():
        try:
            out = convert_file(path, output_dir)
            print(f"    ✅ [{ext}] → {out}")
        except Exception as exc:  # noqa: BLE001
            errors.append(f"{ext}: {exc}")
            print(f"    ❌ [{ext}] 失败：{exc}")

    print("\n=== 总结 ===")
    if errors:
        print("❌ 有问题：")
        for e in errors:
            print("  -", e)
        return 1
    print("✅ 全部转换成功！")
    print(f"输出目录：{output_dir}")
    print("\n输出文件预览：")
    for md in sorted(output_dir.glob("*.md")):
        print(f"\n--- {md.name} ---")
        content = md.read_text(encoding="utf-8").strip().splitlines()
        for line in content[:30]:
            print("  " + line)
        if len(content) > 30:
            print(f"  ... （共 {len(content)} 行）")
    return 0


if __name__ == "__main__":
    sys.exit(main())
